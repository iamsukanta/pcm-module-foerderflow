"""fill_docx_template — port of lib/nachweis/docx-filler.ts.

The monolith fills a user-supplied DOCX via docxtemplater, whose default
delimiters are single braces: placeholders in the template look like
``{PLACEHOLDER_NAME}``. The ``feld_mappings`` map FörderFlow fields to those
placeholders (``{"massnahme_name": "{TITEL}"}``); we strip the braces, resolve
the value from the flat NachweisData snapshot, and substitute ``{TITEL}`` in the
document text. python-docx replaces the placeholders run-by-paragraph (the
monolith's docxtemplater output is plain text substitution as well).
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def _eur(n: float) -> str:
    """de-DE currency-style number: thousands '.', decimal ',', 2 fraction digits."""
    # Format with en-US grouping first, then swap separators to German.
    s = f"{n:,.2f}"  # e.g. "1,234.56" / "-1,234.56"
    return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


def _build_flat_data(data: dict[str, Any]) -> dict[str, str]:
    """Mirror of buildFlatData() in docx-filler.ts."""
    ausgaben = data.get("ausgaben") or []
    gesamt_ausgaben = sum(float(a["betrag_foerderfahig"]) for a in ausgaben)
    einnahmen = data["einnahmen"]
    gesamt_einnahmen = (
        float(einnahmen["eigenmittel"])
        + float(einnahmen["zuwendung"])
        + float(einnahmen["sonstige"])
    )

    generated = datetime.fromisoformat(data["generated_at"])
    date_str = generated.strftime("%d.%m.%Y")

    massnahme = data["massnahme"]
    fy = data["fiscal_year"]
    return {
        "massnahme_name": massnahme["name"],
        "massnahme_foerderquote": f"{float(massnahme['foerderquote']):.2f}",
        "massnahme_budget_gesamt": _eur(float(massnahme["budget_gesamt"])),
        "massnahme_laufzeit_von": massnahme["laufzeit_von"],
        "massnahme_laufzeit_bis": massnahme["laufzeit_bis"],
        "massnahme_funder_name": massnahme["funder_name"],
        "fiscal_year_jahr": str(fy["jahr"]),
        "fiscal_year_beginn": fy["beginn"],
        "fiscal_year_ende": fy["ende"],
        "org_name": data["org"]["name"],
        "eigenmittel": _eur(float(einnahmen["eigenmittel"])),
        "zuwendung": _eur(float(einnahmen["zuwendung"])),
        "sonstige_einnahmen": _eur(float(einnahmen["sonstige"])),
        "gesamt_einnahmen": _eur(gesamt_einnahmen),
        "gesamt_ausgaben": _eur(gesamt_ausgaben),
        "saldo": _eur(gesamt_einnahmen - gesamt_ausgaben),
        "generated_at_datum": date_str,
    }


def _iter_paragraphs(container: Any) -> Any:
    """Yield every paragraph in the document body and (nested) table cells."""
    yield from container.paragraphs
    for table in container.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Any:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _replace_in_paragraph(para: Paragraph, render: dict[str, str]) -> None:
    """Replace ``{KEY}`` placeholders, joining text across runs.

    docxtemplater tolerates placeholders split across runs; python-docx keeps
    them per-run, so we collapse the paragraph text into the first run.
    """
    runs = para.runs
    if not runs:
        return
    full = "".join(run.text for run in runs)
    if not any(key in full for key in render):
        return
    for key, value in render.items():
        full = full.replace(key, value)
    runs[0].text = full
    for run in runs[1:]:
        run.text = ""


def fill_docx_template(
    template_bytes: bytes,
    nachweis_data: dict[str, Any],
    feld_mappings: dict[str, str],
) -> bytes:
    flat = _build_flat_data(nachweis_data)

    # Build {PLACEHOLDER} → value from the field mappings (braces stripped/re-added
    # exactly like the monolith: leading "{" and trailing "}" are optional).
    render: dict[str, str] = {}
    for ff_field, placeholder in feld_mappings.items():
        key = placeholder
        if key.startswith("{"):
            key = key[1:]
        if key.endswith("}"):
            key = key[:-1]
        value = flat.get(ff_field)
        if value is not None:
            render["{" + key + "}"] = value

    doc = Document(io.BytesIO(template_bytes))
    if render:
        for para in _iter_paragraphs(doc):
            _replace_in_paragraph(para, render)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
