"""Verwendungsnachweis + reporting parity tests."""

from __future__ import annotations

import io
import zipfile
from datetime import date
from decimal import Decimal

import pytest
from openpyxl import load_workbook

from app.models.master import FiscalYear
from app.services.ampel import berechne_ampel
from app.services.foerdermassnahme_berechnung import berechne_zuwendung


@pytest.fixture
def base(client, db_session, org):
    fid = client.post("/api/protected/funder", json={"name": "Funder", "typ": "STIFTUNG"}).json()["data"]["id"]
    hjid = client.post(
        "/api/protected/haushaltsjahre",
        json={"jahr": 2026, "beginn": "2026-01-01", "ende": "2026-12-31"},
    ).json()["data"]["id"]
    mid = client.post(
        "/api/protected/foerdermassnahmen",
        json={
            "funder_id": fid, "name": "Maßnahme", "budget_gesamt": 100000, "foerderquote": 80,
            "laufzeit_von": "2026-01-01", "laufzeit_bis": "2026-12-31", "mittelabruf_verfahren": "ABRUF",
        },
    ).json()["data"]["id"]
    return {"funder": fid, "hj": hjid, "measure": mid}


# ── pure calc tests ───────────────────────────────────────────────────────────
def test_berechne_zuwendung():
    assert berechne_zuwendung("ANTEIL", 1000, foerderquote_input=80).zuwendung == 800
    r = berechne_zuwendung("FEHLBEDARF", 1000, eigenmittel=200, drittmittel=100)
    assert r.zuwendung == 700 and r.foerderquote == 70
    assert berechne_zuwendung("FESTBETRAG", 5000).zuwendung == 5000


def test_ampel():
    today = date.today()
    r = berechne_ampel(
        betrag_bewilligt=1000, betrag_ist=980, laufzeit_von=today, laufzeit_bis=today,
        overhead_limit_prozent=None, overhead_ist_prozent=0,
    )
    assert r.status == "ROT"  # >95%
    r = berechne_ampel(
        betrag_bewilligt=1000, betrag_ist=850, laufzeit_von=today, laufzeit_bis=today,
        overhead_limit_prozent=None, overhead_ist_prozent=0,
    )
    assert r.status == "GELB"  # 80-95%


# ── CRUD + transitions ─────────────────────────────────────────────────────────
def test_create_requires_frist_when_no_default(client, base):
    r = client.post(
        "/api/protected/verwendungsnachweise",
        json={
            "funding_measure_id": base["measure"], "fiscal_year_id": base["hj"],
            "zeitraum_von": "2026-01-01", "zeitraum_bis": "2026-12-31", "typ": "VERWENDUNGSNACHWEIS",
        },
    )
    assert r.status_code == 422 and r.json()["code"] == "VALIDATION_FRIST_MISSING"


def test_crud_and_transitions(client, base):
    r = client.post(
        "/api/protected/verwendungsnachweise",
        json={
            "funding_measure_id": base["measure"], "fiscal_year_id": base["hj"],
            "zeitraum_von": "2026-01-01", "zeitraum_bis": "2026-12-31",
            "frist": "2027-04-30", "typ": "VERWENDUNGSNACHWEIS",
        },
    )
    assert r.status_code == 201, r.text
    nid = r.json()["data"]["id"]
    assert r.json()["data"]["status"] == "OFFEN"

    # invalid transition OFFEN -> EINGEREICHT directly via PATCH
    r = client.patch(f"/api/protected/verwendungsnachweise/{nid}", json={"status": "EINGEREICHT"})
    assert r.status_code == 409 and r.json()["code"] == "INVALID_STATUS_TRANSITION"

    # valid OFFEN -> IN_BEARBEITUNG
    r = client.patch(f"/api/protected/verwendungsnachweise/{nid}", json={"status": "IN_BEARBEITUNG"})
    assert r.status_code == 200 and r.json()["data"]["status"] == "IN_BEARBEITUNG"

    # delete only allowed in OFFEN
    r = client.delete(f"/api/protected/verwendungsnachweise/{nid}")
    assert r.status_code == 409 and r.json()["code"] == "INVALID_STATUS"


def test_einreichen_builds_snapshot(client, base):
    nid = client.post(
        "/api/protected/verwendungsnachweise",
        json={
            "funding_measure_id": base["measure"], "fiscal_year_id": base["hj"],
            "zeitraum_von": "2026-01-01", "zeitraum_bis": "2026-12-31",
            "frist": "2027-04-30", "typ": "VERWENDUNGSNACHWEIS",
        },
    ).json()["data"]["id"]

    r = client.post(f"/api/protected/verwendungsnachweise/{nid}/einreichen")
    assert r.status_code == 200, r.text
    d = r.json()["data"]
    assert d["status"] == "EINGEREICHT"
    assert d["snapshot_json"] is not None
    assert d["snapshot_json"]["massnahme"]["name"] == "Maßnahme"
    assert d["eingereicht_am"] is not None

    # snapshot now immutable
    r = client.patch(f"/api/protected/verwendungsnachweise/{nid}", json={"snapshot_json": {}})
    assert r.status_code == 409 and r.json()["code"] == "SNAPSHOT_IMMUTABLE"


def test_preview(client, base):
    r = client.get(f"/api/protected/foerdermassnahmen/{base['measure']}/verwendungsnachweis/preview")
    assert r.status_code == 200, r.text
    d = r.json()["data"]
    assert d["massnahme"]["name"] == "Maßnahme"
    assert d["budget"]["gesamt_ist"] == 0
    assert "ampel_status" in d


def test_soll_ist_position(client, base):
    client.post(
        "/api/protected/finanzplan-positionen",
        json={"funding_measure_id": base["measure"], "positionscode": "1.1",
              "bezeichnung": "Personal", "betrag_bewilligt": 50000},
    )
    r = client.get(f"/api/protected/foerdermassnahmen/{base['measure']}/soll-ist-position")
    assert r.status_code == 200
    payload = r.json()["data"]
    assert len(payload["data"]) == 1
    assert payload["gesamt_bewilligt"] == "50000.00"


def test_excel_report_download(client, base):
    r = client.get(
        f"/api/protected/foerdermassnahmen/{base['measure']}/verwendungsnachweis?fiscal_year_id={base['hj']}"
    )
    assert r.status_code == 200
    # ZIP package (parity with monolith): Excel + belege/ folder, X-Filename hint.
    assert r.headers["content-type"] == "application/zip"
    assert r.headers["x-filename"].endswith(".zip")
    zf = zipfile.ZipFile(io.BytesIO(r.content))
    names = zf.namelist()
    xlsx_names = [n for n in names if n.endswith(".xlsx")]
    assert len(xlsx_names) == 1
    assert any(n.startswith("belege/") for n in names)
    wb = load_workbook(io.BytesIO(zf.read(xlsx_names[0])))
    assert "Einnahmen & Ausgaben" in wb.sheetnames
    assert "Belegliste" in wb.sheetnames
    # title cell
    assert "Verwendungsnachweis 2026" in wb["Einnahmen & Ausgaben"]["A1"].value


def test_excel_report_requires_fy(client, base):
    r = client.get(f"/api/protected/foerdermassnahmen/{base['measure']}/verwendungsnachweis")
    assert r.status_code == 400 and r.json()["code"] == "MISSING_PARAM"


def test_docx_filler_substitutes_placeholders():
    """fill_docx_template replaces {PLACEHOLDER} text via the field mappings."""
    from docx import Document

    from app.services.nachweis.docx_filler import fill_docx_template

    doc = Document()
    doc.add_paragraph("Titel: {TITEL} — Jahr {JAHR}")
    src = io.BytesIO()
    doc.save(src)

    data = {
        "massnahme": {
            "name": "Projekt X",
            "foerderquote": 80.0,
            "budget_gesamt": 100000.0,
            "laufzeit_von": "2026-01-01",
            "laufzeit_bis": "2026-12-31",
            "funder_name": "Funder",
        },
        "fiscal_year": {"jahr": 2026, "beginn": "2026-01-01", "ende": "2026-12-31"},
        "org": {"name": "Org"},
        "einnahmen": {"eigenmittel": 20000.0, "zuwendung": 80000.0, "sonstige": 0},
        "ausgaben": [{"betrag_foerderfahig": 50000.0}],
        "generated_at": "2026-06-18T10:00:00+00:00",
    }
    mappings = {"massnahme_name": "{TITEL}", "fiscal_year_jahr": "{JAHR}"}

    out = fill_docx_template(src.getvalue(), data, mappings)
    result = Document(io.BytesIO(out))
    text = "\n".join(p.text for p in result.paragraphs)
    assert "Projekt X" in text
    assert "2026" in text
    assert "{TITEL}" not in text and "{JAHR}" not in text


def test_stundennachweis_download(client, base):
    r = client.get(
        f"/api/protected/foerdermassnahmen/{base['measure']}/stundennachweis?fiscal_year_id={base['hj']}"
    )
    assert r.status_code == 200
    wb = load_workbook(io.BytesIO(r.content))
    # no payroll data -> placeholder sheet
    assert "Keine Daten" in wb.sheetnames
